# convert_logs_to_docx.py
import os
import json
import tempfile
from docx import Document
from docx.shared import Pt
from pathlib import Path
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def load_inputs():
    """Load the user inputs from JSON file."""
    with open("user_inputs.json", "r") as f:
        inputs = json.load(f)
    return inputs

def convert_md_to_docx(md_path: Path, output_path: Path):
    """Convert markdown content to DOCX with proper formatting"""
    doc = Document()
    
    # Add Nigerian-style document formatting
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    
    # Set document margins
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    current_paragraph = None
    
    for line in content.split("\n"):
        line = line.strip()
        
        # Handle headers
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
            
        # Handle bold sections (Nigerian template headers)
        elif line.startswith("**") and line.endswith("**"):
            current_paragraph = doc.add_paragraph()
            current_paragraph.add_run(line.strip("*")).bold = True
        
        # Handle tables
        elif "|" in line:
            cells = [c.strip() for c in line.split("|") if c.strip()]
            if cells:
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = "Light Shading"
                hdr_cells = table.rows[0].cells
                for i, cell in enumerate(cells):
                    hdr_cells[i].text = cell
                    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Handle image placeholders
        elif "<!-- IMAGE:" in line:
            desc = line.split(":")[1].split("-->")[0].strip()
            paragraph = doc.add_paragraph()
            
            # Create SVG placeholder instead of using file
            svg = f'''
            <svg width="300" height="200" xmlns="http://www.w3.org/2000/svg">
                <rect width="100%" height="100%" fill="#f0f0f0"/>
                <text x="50%" y="50%" dominant-baseline="middle" text-anchor="middle" 
                      font-family="Arial" font-size="14" fill="#666">
                    [IMAGE PLACEHOLDER: {desc}]
                </text>
            </svg>
            '''
            
            # Create temporary SVG file
            with tempfile.NamedTemporaryFile(delete=False, suffix=".svg") as temp_svg:
                temp_svg.write(svg.encode())
                temp_svg_path = temp_svg.name
            
            try:
                run = paragraph.add_run()
                run.add_picture(temp_svg_path, width=Pt(300))
                os.unlink(temp_svg_path)  # Cleanup temporary SVG
            except Exception as e:
                paragraph.add_run(f"[IMAGE PLACEHOLDER: {desc}]").italic = True
            
            paragraph.add_run(f"\nFigure: {desc}").italic = True
            
        # Handle regular content
        else:
            if line:
                if current_paragraph:
                    current_paragraph.add_run("\n" + line)
                else:
                    current_paragraph = doc.add_paragraph(line)
            else:
                current_paragraph = None

    doc.save(output_path)